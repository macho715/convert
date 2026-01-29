#!/usr/bin/env python3
"""
템플릿 내 Placeholder를 JSON 값으로 치환
"""

import argparse
import json
from docx import Document


def _flatten_value(val):
    if isinstance(val, list):
        return "\n\n".join(str(x) for x in val)
    return str(val) if val is not None else ""


def replace_in_paragraph(par, mapping: dict) -> None:
    for run in par.runs:
        text = run.text
        for key, val in mapping.items():
            token = f"{{{{{key}}}}}"
            if token in text:
                run.text = text.replace(token, _flatten_value(val))


def fill_docx(template_path: str, content_json: str, out_path: str) -> None:
    doc = Document(template_path)
    with open(content_json, "r", encoding="utf-8") as f:
        data = json.load(f)

    for section in doc.sections:
        for par in section.header.paragraphs:
            replace_in_paragraph(par, data)

    for par in doc.paragraphs:
        replace_in_paragraph(par, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for par in cell.paragraphs:
                    replace_in_paragraph(par, data)

    doc.save(out_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="템플릿 콘텐츠 채움")
    parser.add_argument("--template", "-t", required=True)
    parser.add_argument("--content-json", "-c", required=True)
    parser.add_argument("--out", "-o", required=True)
    args = parser.parse_args()
    fill_docx(args.template, args.content_json, args.out)
