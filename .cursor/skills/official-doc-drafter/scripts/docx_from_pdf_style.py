#!/usr/bin/env python3
"""
PDF 스타일 기반 Word 문서 생성 (본문/헤더 구조 포함)
"""

import argparse
import json
from docx import Document
from docx.shared import Mm, Pt


def create_from_style(style_json: str, content_json: str, out_docx: str) -> None:
    with open(style_json, "r", encoding="utf-8") as f:
        style = json.load(f)
    with open(content_json, "r", encoding="utf-8") as f:
        content = json.load(f)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    marg = style.get("margins_mm", {})
    sec.left_margin = Mm(marg.get("left", 25))
    sec.right_margin = Mm(marg.get("right", 25))
    sec.top_margin = Mm(25)
    sec.bottom_margin = Mm(25)

    body_font = style.get("body_font", {})
    normal_style = doc.styles["Normal"].font
    normal_style.name = body_font.get("name", "Malgun Gothic")
    normal_style.size = Pt(body_font.get("size_pt", 11))

    header = sec.header
    header_para = header.add_paragraph()
    header_para.add_run("{{LOGO}} ")
    header_para.add_run(content.get("ORG_NAME", "")).bold = True

    doc.add_paragraph(f"문서번호: {content.get('DOC_NO', '')}")
    doc.add_paragraph(f"일자: {content.get('DATE', '')}")
    doc.add_paragraph(f"수신: {content.get('RECIPIENT', '')}")
    doc.add_paragraph(content.get("SUBJECT", ""))
    body = content.get("BODY_PARAGRAPHS", [])
    if isinstance(body, str):
        body = [body]
    for p in body:
        doc.add_paragraph(p)

    doc.save(out_docx)


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--style-json", "-s", required=True)
    p.add_argument("--content-json", "-c", required=True)
    p.add_argument("--out", "-o", required=True)
    args = p.parse_args()
    create_from_style(args.style_json, args.content_json, args.out)
