#!/usr/bin/env python3
"""
Word 헤더에 로고 삽입
"""

import argparse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Mm


def insert_logo(
    docx_path: str, logo_path: str, out_path: str, width_mm: float | None = None
) -> None:
    doc = Document(docx_path)
    sec = doc.sections[0]
    header = sec.header
    for par in header.paragraphs:
        if "{{LOGO}}" in par.text:
            par.clear()
            run = par.add_run()
            if width_mm is not None:
                run.add_picture(logo_path, width=Mm(width_mm))
            else:
                run.add_picture(logo_path)
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.save(out_path)


if __name__ == "__main__":
    p = argparse.ArgumentParser()
    p.add_argument("--docx", "-d", required=True)
    p.add_argument("--logo", "-l", required=True)
    p.add_argument("--out", "-o", required=True)
    p.add_argument("--width-mm", "-w", type=float, default=None)
    args = p.parse_args()
    insert_logo(args.docx, args.logo, args.out, args.width_mm)
