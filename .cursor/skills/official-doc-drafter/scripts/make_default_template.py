#!/usr/bin/env python3
"""
A4 기반 기본 공문서 Word 템플릿 생성
"""

import argparse
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_cell_border(cell, **kwargs):
    """
    Table cell 모든 테두리 설정
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge, value in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in value.items():
            el.set(qn(f"w:{k}"), str(v))
        tcPr.append(el)


def create_template(out_path: str, org_name: str = "{{ORG_NAME}}") -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Mm(25))

    header = section.header
    hdr_table = header.add_table(rows=1, cols=2)
    hdr_table.autofit = True
    left, right = hdr_table.rows[0].cells
    p_logo = left.paragraphs[0]
    run_logo = p_logo.add_run("{{LOGO}}")
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_org = right.paragraphs[0]
    p_org.add_run(org_name).bold = True
    p_org.add_run("\n{{ORG_ADDRESS}}")
    p_org.add_run("\n{{ORG_CONTACT}}")
    set_cell_border(
        right, bottom={"val": "single", "sz": "6", "space": "0", "color": "000000"}
    )

    doc.add_paragraph()

    doc.add_paragraph("문서번호: {{DOC_NO}}")
    doc.add_paragraph("일자: {{DATE}}")
    doc.add_paragraph("수신: {{RECIPIENT}}")
    doc.add_paragraph("제목: {{SUBJECT}}").runs[0].font.size = Pt(14)

    doc.add_paragraph("{{BODY_PARAGRAPHS}}")

    doc.add_paragraph("첨부: {{ATTACHMENTS}}")

    doc.add_paragraph("{{SIGNER_TITLE}}\n{{SIGNER_NAME}}")

    doc.save(out_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="기본 공문서 템플릿 생성")
    parser.add_argument("--out", "-o", required=True, help="출력 .docx 경로")
    parser.add_argument("--org-name", "-n", default="{{ORG_NAME}}", help="기관명")
    args = parser.parse_args()
    create_template(args.out, args.org_name)
