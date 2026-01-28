"""
HTML to Word Document Converter
Converts UAE_HVDC_GreatWaters_SPMT_Loadout_Approval.html to DOCX format
with proper page settings matching the image specifications
"""
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def html_to_word(html_path: str, docx_path: str = None) -> str:
    """
    Convert HTML file to Word document with proper formatting
    
    Args:
        html_path: Path to input HTML file
        docx_path: Path to output DOCX file (optional)
    
    Returns:
        Path to created DOCX file
    """
    html_file = Path(html_path)
    if not html_file.exists():
        raise FileNotFoundError(f"HTML file not found: {html_path}")
    
    # Generate output path if not provided
    if docx_path is None:
        docx_path = html_file.with_suffix('.docx')
    else:
        docx_path = Path(docx_path)
    
    # Read HTML content
    with open(html_file, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    # Parse HTML
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create Word document
    doc = Document()
    
    # Set page margins (from image: Top 3cm, Bottom 2.75cm, Left 2cm, Right 2.25cm)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.75)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.25)
    
    # Set default font (BatangChe for Korean compatibility)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'BatangChe'
    font.size = Pt(11)
    
    # Process body content
    body = soup.find('body')
    if body:
        _process_html_body(doc, body)
    else:
        _process_html_body(doc, soup)
    
    # Save document
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    
    print(f"✓ Successfully converted: {html_file.name} → {docx_path.name}")
    print(f"  Output location: {docx_path.absolute()}")
    return str(docx_path)


def _process_html_body(doc: Document, body) -> None:
    """Process HTML body and convert to Word document"""
    
    # Find all page divs
    pages = body.find_all('div', class_='page')
    
    for page_idx, page in enumerate(pages):
        if page_idx > 0:
            # Add page break for subsequent pages
            doc.add_page_break()
        
        # Process all elements in the page
        for element in page.children:
            if hasattr(element, 'name') and element.name is not None:
                tag_name = element.name.lower()
                
                # Title paragraphs
                if tag_name == 'p' and 'title' in element.get('class', []):
                    text = element.get_text(strip=True)
                    if text:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(text)
                        run.bold = True
                        run.font.size = Pt(12)
                        run.underline = True
                
                # Date paragraph
                elif tag_name == 'p' and 'date' in element.get('class', []):
                    text = element.get_text(strip=True)
                    if text:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run = p.add_run(text)
                        run.font.size = Pt(10)
                        p.paragraph_format.space_after = Pt(12)
                
                # Box table (Project details)
                elif tag_name == 'table' and 'box' in element.get('class', []):
                    _process_box_table(doc, element)
                
                # Grid tables
                elif tag_name == 'table' and 'grid' in element.get('class', []):
                    _process_grid_table(doc, element)
                
                # Section headings
                elif tag_name == 'div' and 'section' in element.get('class', []):
                    text = element.get_text(strip=True)
                    # Remove the □ symbol if present (we'll add it back)
                    text = text.replace('□', '').strip()
                    if text:
                        p = doc.add_paragraph()
                        run = p.add_run(f'□ {text}')
                        run.bold = True
                        run.font.size = Pt(11)
                        p.paragraph_format.space_before = Pt(14)
                        p.paragraph_format.space_after = Pt(6)
                
                # Regular paragraphs
                elif tag_name == 'p':
                    classes = element.get('class', [])
                    text = element.get_text(strip=True)
                    
                    if text:
                        p = doc.add_paragraph()
                        
                        # Apply indentation
                        if 'indent' in classes:
                            p.paragraph_format.left_indent = Inches(0.25)
                        
                        # Apply alignment
                        if 'end' in classes:
                            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            p.paragraph_format.space_before = Pt(16)
                        
                        # Add text with formatting
                        _add_text_with_formatting(p, element)
                        p.paragraph_format.space_after = Pt(6)
                
                # Page number
                elif tag_name == 'div' and 'pagenum' in element.get('class', []):
                    text = element.get_text(strip=True)
                    if text:
                        p = doc.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run(text)
                        run.font.size = Pt(10)
                        p.paragraph_format.space_before = Pt(18)


def _process_box_table(doc: Document, table_element) -> None:
    """Process box table (Project details)"""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Create a single-cell table for the box
    word_table = doc.add_table(rows=1, cols=1)
    word_table.style = 'Table Grid'
    
    cell = word_table.rows[0].cells[0]
    
    # Get all divs from the table cell
    for row in rows:
        for td in row.find_all('td'):
            for div in td.find_all('div'):
                text = div.get_text(strip=True)
                if text:
                    p = cell.add_paragraph(text)
                    # Make Vendor and Amount bold
                    if 'Vendor' in text or 'Amount' in text:
                        for run in p.runs:
                            run.bold = True
    
    # Set cell padding
    cell.vertical_alignment = 1  # Top alignment
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)


def _process_grid_table(doc: Document, table_element) -> None:
    """Process grid table with proper formatting"""
    rows = table_element.find_all('tr')
    if not rows:
        return
    
    # Determine number of columns
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(['td', 'th']))
        max_cols = max(max_cols, cols)
    
    if max_cols == 0:
        return
    
    # Create Word table
    word_table = doc.add_table(rows=len(rows), cols=max_cols)
    word_table.style = 'Light Grid Accent 1'
    
    # Populate table
    for i, row in enumerate(rows):
        cells = row.find_all(['td', 'th'])
        for j, cell in enumerate(cells):
            if j < max_cols:
                word_cell = word_table.rows[i].cells[j]
                text = cell.get_text(strip=True)
                word_cell.text = text
                
                # Style header cells
                if cell.name == 'th' or i == 0:
                    for paragraph in word_cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0, 0, 0)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Apply cell alignment
                else:
                    # Check if it's a number (for right alignment)
                    try:
                        float(text.replace(',', ''))
                        word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    except ValueError:
                        # First column left, others center
                        if j == 0:
                            word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        else:
                            word_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_text_with_formatting(paragraph, element) -> None:
    """Add text with inline formatting (bold, italic, etc.)"""
    # Get all text nodes and inline elements
    for content in element.descendants:
        if isinstance(content, str):
            if content.strip():
                paragraph.add_run(content)
        elif hasattr(content, 'name'):
            tag = content.name.lower()
            text = content.get_text(strip=True)
            
            if tag == 'b' or tag == 'strong':
                run = paragraph.add_run(text)
                run.bold = True
            elif tag == 'i' or tag == 'em':
                run = paragraph.add_run(text)
                run.italic = True
            elif tag == 'u':
                run = paragraph.add_run(text)
                run.underline = True
            elif tag == 'br':
                paragraph.add_run('\n')


if __name__ == '__main__':
    import sys
    
    # Default input file
    html_file = r"c:\Users\SAMSUNG\Desktop\UAE_HVDC_GreatWaters_SPMT_Loadout_Approval.html"
    
    if len(sys.argv) > 1:
        html_file = sys.argv[1]
    
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        result = html_to_word(html_file, output_file)
        print(f"\n✓ Conversion complete!")
        print(f"  Word document saved to: {result}")
    except Exception as e:
        print(f"✗ Error: {e}")
        import traceback
        traceback.print_exc()

