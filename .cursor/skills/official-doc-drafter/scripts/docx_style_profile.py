#!/usr/bin/env python3
"""
Create a compact, heuristic style profile for a DOCX file.
This is NOT a pixel renderer. It extracts what python-docx can observe:
- page size / margins per section
- dominant font name / size distribution (runs)
- paragraph alignment / spacing / line spacing distribution
- leftover placeholders like {{TOKEN}}
"""

from __future__ import annotations

import argparse
import json
import re
from collections import Counter
from typing import Any, Dict, Optional, Tuple

from docx import Document


PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")


def _len_mm(x) -> Optional[float]:
    try:
        return float(x.mm) if x is not None else None
    except Exception:
        return None


def _len_pt(x) -> Optional[float]:
    try:
        return float(x.pt) if x is not None else None
    except Exception:
        return None


def _alignment_name(alignment) -> Optional[str]:
    if alignment is None:
        return None
    # python-docx enums stringify like 'CENTER (1)'
    s = str(alignment)
    return s.split()[0] if s else None


def _effective_font(run, par, doc) -> Tuple[Optional[str], Optional[float]]:
    """
    Return (font_name, font_size_pt) using fallback chain:
    run -> paragraph style -> Normal style
    """
    name = None
    size_pt = None

    try:
        if run is not None and run.font is not None:
            if run.font.name:
                name = run.font.name
            if run.font.size:
                size_pt = _len_pt(run.font.size)
    except Exception:
        pass

    try:
        if (
            (name is None or size_pt is None)
            and par is not None
            and par.style is not None
        ):
            f = par.style.font
            if name is None and f is not None and f.name:
                name = f.name
            if size_pt is None and f is not None and f.size:
                size_pt = _len_pt(f.size)
    except Exception:
        pass

    try:
        if (name is None or size_pt is None) and doc is not None:
            f = doc.styles["Normal"].font
            if name is None and f is not None and f.name:
                name = f.name
            if size_pt is None and f is not None and f.size:
                size_pt = _len_pt(f.size)
    except Exception:
        pass

    return name, size_pt


def _line_spacing_repr(par) -> Optional[Dict[str, Any]]:
    """
    python-docx returns either:
      - None
      - float (multiple)
      - Length (absolute)
    """
    try:
        ls = par.paragraph_format.line_spacing
    except Exception:
        return None

    if ls is None:
        return None
    if isinstance(ls, (int, float)):
        return {"type": "multiple", "value": float(ls)}
    # assume Length-like
    pt = _len_pt(ls)
    return {"type": "exact_pt", "value": pt}


def _scan_placeholders(doc: Document) -> Counter:
    tokens = Counter()

    def scan_text(text: str):
        for m in PLACEHOLDER_RE.findall(text or ""):
            tokens[m] += 1

    # body paragraphs
    for p in doc.paragraphs:
        scan_text(p.text)

    # tables
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan_text(p.text)

    # headers/footers
    for sec in doc.sections:
        for p in sec.header.paragraphs:
            scan_text(p.text)
        for p in sec.footer.paragraphs:
            scan_text(p.text)

    return tokens


def profile_docx(docx_path: str, max_paragraphs: int = 400) -> Dict[str, Any]:
    doc = Document(docx_path)

    # sections
    sections = []
    for sec in doc.sections:
        sections.append(
            {
                "page_width_mm": _len_mm(sec.page_width),
                "page_height_mm": _len_mm(sec.page_height),
                "margins_mm": {
                    "top": _len_mm(sec.top_margin),
                    "bottom": _len_mm(sec.bottom_margin),
                    "left": _len_mm(sec.left_margin),
                    "right": _len_mm(sec.right_margin),
                },
                "header_distance_mm": _len_mm(getattr(sec, "header_distance", None)),
                "footer_distance_mm": _len_mm(getattr(sec, "footer_distance", None)),
            }
        )

    # Normal style (defaults)
    normal_font = {"name": None, "size_pt": None}
    try:
        nf = doc.styles["Normal"].font
        normal_font["name"] = nf.name
        normal_font["size_pt"] = _len_pt(nf.size)
    except Exception:
        pass

    font_names = Counter()
    font_sizes = Counter()
    alignments = Counter()
    style_names = Counter()
    space_before_pt = Counter()
    space_after_pt = Counter()
    line_spacing = Counter()

    # paragraph sampling
    for i, p in enumerate(doc.paragraphs[:max_paragraphs]):
        if p.text is None or not p.text.strip():
            continue

        style_names[str(getattr(p.style, "name", None))] += 1
        alignments[_alignment_name(p.alignment)] += 1

        # spacing
        try:
            fmt = p.paragraph_format
            sb = _len_pt(fmt.space_before)
            sa = _len_pt(fmt.space_after)
            if sb is not None:
                space_before_pt[round(sb, 1)] += 1
            if sa is not None:
                space_after_pt[round(sa, 1)] += 1
            ls = _line_spacing_repr(p)
            if ls is not None:
                key = f"{ls.get('type')}:{ls.get('value')}"
                line_spacing[key] += 1
        except Exception:
            pass

        # run-level font sampling
        for r in p.runs:
            name, size_pt = _effective_font(r, p, doc)
            if name:
                font_names[name] += 1
            if size_pt is not None:
                font_sizes[round(size_pt, 1)] += 1

    placeholders = _scan_placeholders(doc)

    dominant_font = font_names.most_common(1)[0][0] if font_names else None
    dominant_size = font_sizes.most_common(1)[0][0] if font_sizes else None

    return {
        "docx_path": docx_path,
        "sections": sections,
        "normal_font": normal_font,
        "dominant_font": dominant_font,
        "dominant_font_size_pt": dominant_size,
        "counters": {
            "font_names_top": font_names.most_common(15),
            "font_sizes_top": font_sizes.most_common(15),
            "alignments_top": alignments.most_common(10),
            "style_names_top": style_names.most_common(15),
            "space_before_pt_top": space_before_pt.most_common(10),
            "space_after_pt_top": space_after_pt.most_common(10),
            "line_spacing_top": line_spacing.most_common(10),
        },
        "placeholders": placeholders.most_common(50),
    }


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--docx", required=True, help="Input .docx path")
    ap.add_argument("--out-json", required=True, help="Output JSON path")
    ap.add_argument("--max-paragraphs", type=int, default=400)
    args = ap.parse_args()

    prof = profile_docx(args.docx, max_paragraphs=args.max_paragraphs)
    with open(args.out_json, "w", encoding="utf-8") as f:
        json.dump(prof, f, ensure_ascii=False, indent=2)

    print(f"Wrote profile: {args.out_json}")


if __name__ == "__main__":
    main()
