#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Excel 파일을 Word 문서로 변환 (포맷, 스타일, 함수 유지)

Excel 파일의 모든 시트를 확인하여 포맷, 스타일, 함수를 그대로 유지하면서
Word 문서로 변환합니다.
"""

import sys
from pathlib import Path
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

try:
    from openpyxl.cell.cell import Cell
    from openpyxl.cell.read_only import ReadOnlyCell
except ImportError:
    pass


def excel_color_to_rgb(color):
    """Excel 색상을 RGB로 변환"""
    if color is None:
        return None
    if hasattr(color, 'rgb'):
        rgb = color.rgb
        if rgb:
            # 'FFRRGGBB' 형식에서 RGB 추출
            if isinstance(rgb, str) and len(rgb) == 8:
                r = int(rgb[2:4], 16)
                g = int(rgb[4:6], 16)
                b = int(rgb[6:8], 16)
                return RGBColor(r, g, b)
    return None


def get_cell_value(cell):
    """셀 값 가져오기 (함수면 함수 반환, 아니면 값 반환)"""
    if cell.data_type == 'f':  # formula
        return f'={cell.value}' if cell.value else ''
    if cell.value is None:
        return ''
    return str(cell.value)


def apply_cell_formatting(paragraph, cell, run):
    """셀 포맷을 Word 텍스트에 적용"""
    # 폰트 스타일
    if cell.font:
        if cell.font.bold:
            run.bold = True
        if cell.font.italic:
            run.italic = True
        if cell.font.underline:
            run.underline = True
        if cell.font.size:
            run.font.size = Pt(cell.font.size)
        if cell.font.color:
            rgb = excel_color_to_rgb(cell.font.color)
            if rgb:
                run.font.color = rgb
    
    # 배경색
    if cell.fill and cell.fill.patternType == 'solid':
        if hasattr(cell.fill, 'fgColor'):
            rgb = excel_color_to_rgb(cell.fill.fgColor)
            if rgb:
                # Word에서는 highlight 사용 불가하므로 주석으로 표시
                pass
    
    # 정렬
    if cell.alignment:
        if cell.alignment.horizontal == 'center':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif cell.alignment.horizontal == 'right':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif cell.alignment.horizontal == 'left':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def excel_to_word(excel_path: str, word_path: str = None):
    """
    Excel 파일을 Word 문서로 변환
    
    Args:
        excel_path: Excel 파일 경로
        word_path: 출력 Word 파일 경로 (None이면 자동 생성)
    """
    excel_path = Path(excel_path)
    if not excel_path.exists():
        raise FileNotFoundError(f"Excel file not found: {excel_path}")
    
    if word_path is None:
        word_path = excel_path.with_suffix('.docx')
    else:
        word_path = Path(word_path)
    
    print(f"Reading Excel file: {excel_path}")
    wb = load_workbook(excel_path, data_only=False)  # data_only=False로 함수 유지
    
    print(f"  Sheets: {len(wb.sheetnames)}")
    for sheet_name in wb.sheetnames:
        print(f"    - {sheet_name}")
    
    # Word 문서 생성
    doc = Document()
    doc.core_properties.title = excel_path.stem
    doc.core_properties.author = "Excel to Word Converter"
    
    # 각 시트 처리
    for sheet_idx, sheet_name in enumerate(wb.sheetnames):
        ws = wb[sheet_name]
        
        print(f"\nProcessing sheet: {sheet_name} ({ws.max_row} rows x {ws.max_column} cols)")
        
        # 시트 제목 추가
        if sheet_idx > 0:
            doc.add_page_break()
        
        heading = doc.add_heading(sheet_name, level=1)
        
        # 테이블 생성
        # 먼저 데이터 범위 확인
        max_row = ws.max_row
        max_col = ws.max_column
        
        if max_row == 0 or max_col == 0:
            doc.add_paragraph("(Empty sheet)")
            continue
        
        # 테이블 생성
        table = doc.add_table(rows=max_row, cols=max_col)
        table.style = 'Light Grid Accent 1'
        
        # 각 셀 처리
        for row_idx, row in enumerate(ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col), 1):
            for col_idx, cell in enumerate(row, 1):
                word_cell = table.rows[row_idx - 1].cells[col_idx - 1]
                
                # 값 설정
                cell_value = get_cell_value(cell)
                
                # Word 셀에 텍스트 추가
                if cell_value:
                    paragraph = word_cell.paragraphs[0]
                    run = paragraph.add_run(cell_value)
                    
                    # 포맷 적용
                    apply_cell_formatting(paragraph, cell, run)
                
                # 셀 너비 조정 (Excel 열 너비 기반)
                col_letter = get_column_letter(col_idx)
                if col_letter in ws.column_dimensions:
                    col_dim = ws.column_dimensions[col_letter]
                    if col_dim.width:
                        col_width = col_dim.width
                        # Word에서는 인치 단위로 변환 (Excel 기본 단위는 문자 너비)
                        word_cell.width = Inches(col_width * 0.15)  # 대략적인 변환
        
        print(f"  Table created: {max_row} rows x {max_col} cols")
        
        # 함수 정보 추가 (함수가 있는 셀 목록)
        formulas = []
        for row in ws.iter_rows():
            for cell in row:
                if cell.data_type == 'f' and cell.value:
                    formulas.append({
                        'cell': cell.coordinate,
                        'formula': cell.value
                    })
        
        if formulas:
            doc.add_paragraph()  # 빈 줄
            doc.add_heading('Formulas', level=2)
            for f in formulas[:50]:  # 최대 50개만 표시
                p = doc.add_paragraph(f"{f['cell']}: {f['formula']}", style='List Bullet')
    
    # 문서 저장
    print(f"\nSaving Word document: {word_path}")
    doc.save(str(word_path))
    print(f"  Conversion complete!")
    
    return word_path


def main():
    """메인 함수"""
    if len(sys.argv) < 2:
        print("Usage: python excel_to_word_converter.py <excel_file.xlsx> [output.docx]")
        print("\nExample:")
        print("  python excel_to_word_converter.py content-calendar.xlsx")
        print("  python excel_to_word_converter.py content-calendar.xlsx output.docx")
        sys.exit(1)
    
    excel_path = sys.argv[1]
    word_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        output_path = excel_to_word(excel_path, word_path)
        print(f"\nSuccess: {output_path}")
    except Exception as e:
        print(f"\nError: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()

