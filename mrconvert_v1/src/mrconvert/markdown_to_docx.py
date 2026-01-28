from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from .types import ConversionResult, ConversionError


def extract_json_ld_from_markdown(content: str) -> Optional[Dict[str, Any]]:
    """Extract JSON-LD metadata from markdown file"""
    # JSON code block 패턴 찾기
    json_pattern = r"```json\s*(\{.*?\})\s*```"
    match = re.search(json_pattern, content, re.DOTALL)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            return None
    return None


def markdown_to_docx(
    src: Path, 
    dst: Path,
    preserve_metadata: bool = True
) -> ConversionResult:
    """
    Convert markdown file to DOCX format
    
    Args:
        src: Source markdown file path
        dst: Destination DOCX file path
        preserve_metadata: Whether to include JSON-LD metadata in document
        
    Returns:
        ConversionResult with conversion details
    """
    if not src.exists():
        raise FileNotFoundError(f"Source file not found: {src}")
    
    # Read markdown content
    content = src.read_text(encoding='utf-8')
    
    # Extract JSON-LD metadata if present
    metadata = None
    if preserve_metadata:
        metadata = extract_json_ld_from_markdown(content)
        # Remove JSON block from markdown content for cleaner processing
        if metadata:
            content = re.sub(r"```json\s*\{.*?\}\s*```", "", content, flags=re.DOTALL)
    
    # Create Word document
    doc = Document()
    
    # Set default font and styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Enhance heading styles
    for level in range(1, 10):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        heading_style.font.bold = True
        if level == 1:
            heading_style.font.size = Pt(18)
            heading_style.font.color.rgb = RGBColor(0, 32, 96)  # Dark blue
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = RGBColor(0, 51, 102)
        elif level == 3:
            heading_style.font.size = Pt(14)
            heading_style.font.color.rgb = RGBColor(0, 70, 127)
    
    # Add metadata section if available
    if metadata:
        _add_metadata_section(doc, metadata)
    
    # Parse and convert markdown content
    _parse_markdown_to_docx(doc, content)
    
    # Save document
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    
    if not dst.exists():
        raise ConversionError(f"Failed to create DOCX file: {dst}")
    
    return ConversionResult(src, dst, "markdown-to-docx")


def _add_metadata_section(doc: Document, metadata: Dict[str, Any]) -> None:
    """Add JSON-LD metadata as a formatted section in the document"""
    # Title
    title = doc.add_heading('Email Thread Metadata', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subject
    if 'subject' in metadata:
        p = doc.add_paragraph()
        run = p.add_run('Subject: ')
        run.bold = True
        p.add_run(metadata['subject'])
    
    # Date Range
    if 'dateRange' in metadata:
        p = doc.add_paragraph()
        run = p.add_run('Date Range: ')
        run.bold = True
        date_range = metadata['dateRange']
        date_str = f"{date_range.get('start', 'N/A')} to {date_range.get('end', 'N/A')}"
        p.add_run(date_str)
    
    # Participants
    if 'participants' in metadata and metadata['participants']:
        doc.add_paragraph()
        heading = doc.add_heading('Participants', level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = 'Name'
        header_cells[1].text = 'Organization'
        header_cells[2].text = 'Email'
        for cell in header_cells:
            cell.paragraphs[0].runs[0].bold = True
        
        # Data rows
        for participant in metadata['participants']:
            row = table.add_row()
            row.cells[0].text = participant.get('name', 'N/A')
            row.cells[1].text = participant.get('org', 'N/A')
            row.cells[2].text = participant.get('email', 'N/A')
    
    # Topics
    if 'topics' in metadata and metadata['topics']:
        doc.add_paragraph()
        heading = doc.add_heading('Topics', level=2)
        p = doc.add_paragraph(', '.join(metadata['topics']))
    
    # Actions
    if 'actions' in metadata and metadata['actions']:
        doc.add_paragraph()
        heading = doc.add_heading('Actions', level=2)
        for action in metadata['actions']:
            p = doc.add_paragraph(action.get('action', ''), style='List Bullet')
            owner = action.get('owner', '')
            if owner:
                run = p.add_run(f' ({owner})')
                run.italic = True
    
    # Issues
    if 'issues' in metadata and metadata['issues']:
        doc.add_paragraph()
        heading = doc.add_heading('Issues', level=2)
        for issue in metadata['issues']:
            p = doc.add_paragraph(issue.get('description', ''), style='List Bullet')
            issue_type = issue.get('type', '')
            if issue_type:
                run = p.add_run(f' [Type: {issue_type}]')
                run.italic = True
    
    # Add page break before main content
    doc.add_page_break()


def _parse_markdown_to_docx(doc: Document, content: str) -> None:
    """Parse markdown content and add to Word document"""
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines at start
        if not line and i == 0:
            i += 1
            continue
        
        # Email message header (#### Msg X — Name @ Date)
        if re.match(r'^####\s+Msg\s+\d+', line):
            msg_info = _parse_email_header(line)
            if msg_info:
                _add_email_message_header(doc, msg_info)
                # Look for email metadata table and content
                i += 1
                # Skip to next non-empty line
                while i < len(lines) and not lines[i].strip():
                    i += 1
                # Process email metadata table if present
                if i < len(lines) and '|' in lines[i] and 'Key' in lines[i]:
                    table_data = _parse_table(lines, i)
                    if table_data:
                        _add_email_metadata_table(doc, table_data)
                        i += len(table_data) + 1  # +1 for separator row
                # Process email body (code block)
                if i < len(lines) and lines[i].strip().startswith('```'):
                    code_block = _parse_code_block(lines, i)
                    if code_block:
                        _add_email_body(doc, code_block)
                        i += len(code_block.split('\n')) + 2  # +2 for ``` markers
                        continue
            else:
                # Fallback to regular heading
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if text:
                    doc.add_heading(text, level=min(level, 9))
        
        # Regular headings
        elif line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if text:
                doc.add_heading(text, level=min(level, 9))
        
        # Horizontal rule
        elif line.strip() == '---':
            doc.add_paragraph('─' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Table detection (simple markdown table)
        elif '|' in line and line.strip().startswith('|'):
            table_data = _parse_table(lines, i)
            if table_data:
                _add_table_to_docx(doc, table_data)
                # Skip processed lines
                i += len(table_data) - 1
        
        # Code block
        elif line.strip().startswith('```'):
            code_block = _parse_code_block(lines, i)
            if code_block:
                _add_code_block_to_docx(doc, code_block)
                # Skip processed lines
                i += len(code_block.split('\n')) + 1
                continue
        
        # List items
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            list_items = _parse_list(lines, i)
            if list_items:
                _add_list_to_docx(doc, list_items)
                i += len(list_items) - 1
        
        # Regular paragraph
        elif line.strip():
            _add_paragraph_with_formatting(doc, line)
        
        i += 1


def _parse_table(lines: List[str], start_idx: int) -> List[List[str]]:
    """Parse markdown table into list of rows"""
    table_data = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        
        # Skip separator row (|---|---|)
        if re.match(r'^\|[\s\-:]+\|', line):
            i += 1
            continue
        
        # Parse row
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            table_data.append(cells)
        
        i += 1
    
    return table_data if len(table_data) > 1 else []


def _add_table_to_docx(doc: Document, table_data: List[List[str]]) -> None:
    """Add table to Word document with enhanced styling"""
    if not table_data:
        return
    
    table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
    table.style = 'Light Grid Accent 1'
    
    for i, row_data in enumerate(table_data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # Style header row - make text bold
            if i == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Black text for better readability


def _parse_code_block(lines: List[str], start_idx: int) -> Optional[str]:
    """Parse code block from markdown"""
    if start_idx >= len(lines):
        return None
    
    first_line = lines[start_idx].strip()
    if not first_line.startswith('```'):
        return None
    
    # Get language if specified
    language = first_line[3:].strip() if len(first_line) > 3 else ''
    
    code_lines = []
    i = start_idx + 1
    
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('```'):
            break
        code_lines.append(line)
        i += 1
    
    return '\n'.join(code_lines)


def _add_code_block_to_docx(doc: Document, code: str) -> None:
    """Add code block to Word document with monospace font"""
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
    
    # Add light gray background (simulated with shading)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def _parse_list(lines: List[str], start_idx: int) -> List[str]:
    """Parse list items from markdown"""
    items = []
    i = start_idx
    
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith('- ') or line.startswith('* '):
            items.append(line[2:].strip())
        elif line.startswith(' ') and items:
            # Continuation of previous item
            items[-1] += ' ' + line.strip()
        else:
            break
        i += 1
    
    return items


def _add_list_to_docx(doc: Document, items: List[str]) -> None:
    """Add list to Word document"""
    for item in items:
        # Remove markdown formatting from list item
        cleaned_item = _clean_markdown_formatting(item)
        doc.add_paragraph(cleaned_item, style='List Bullet')


def _clean_markdown_formatting(text: str) -> str:
    """Remove markdown formatting from text"""
    # Remove bold/italic markers
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'__(.*?)__', r'\1', text)
    text = re.sub(r'_(.*?)_', r'\1', text)
    
    # Remove links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    return text


def _parse_email_header(line: str) -> Optional[Dict[str, str]]:
    """Parse email message header line: '#### Msg 13 — Name @ 2025-09-19 15:13 +04:00 {#msg-13}'"""
    pattern = r'^####\s+Msg\s+(\d+)\s+—\s+(.+?)\s+@\s+(.+?)\s+\{#(.+?)\}'
    match = re.match(pattern, line)
    if match:
        return {
            'msg_num': match.group(1),
            'sender': match.group(2).strip(),
            'date': match.group(3).strip(),
            'anchor': match.group(4)
        }
    return None


def _add_email_message_header(doc: Document, msg_info: Dict[str, str]) -> None:
    """Add email message header with styling"""
    heading_text = f"Message {msg_info['msg_num']} — {msg_info['sender']} @ {msg_info['date']}"
    heading = doc.add_heading(heading_text, level=4)
    # Apply styling to heading
    for run in heading.runs:
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue


def _add_email_metadata_table(doc: Document, table_data: List[List[str]]) -> None:
    """Add email metadata table with special styling"""
    if not table_data or len(table_data) < 2:
        return
    
    table = doc.add_table(rows=len(table_data) - 1, cols=2)  # Skip header row
    table.style = 'Light List Accent 1'
    
    # Process data rows (skip header row)
    for i, row_data in enumerate(table_data[1:], 0):
        if len(row_data) >= 2:
            table.rows[i].cells[0].text = row_data[0]
            table.rows[i].cells[1].text = row_data[1]
            # Make key column bold
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True


def _add_email_body(doc: Document, body_text: str) -> None:
    """Add email body text with formatting"""
    p = doc.add_paragraph()
    p.style = 'Normal'
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.25)
    
    run = p.add_run(body_text.strip())
    run.font.size = Pt(10)
    run.font.name = 'Calibri'


def _add_paragraph_with_formatting(doc: Document, text: str) -> None:
    """Add paragraph with markdown formatting preserved"""
    cleaned = _clean_markdown_formatting(text)
    
    # Check for bold/italic patterns
    p = doc.add_paragraph()
    
    # Simple inline formatting support
    parts = re.split(r'(\*\*.*?\*\*|_.*?_|`.*?`)', cleaned)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('_') and part.endswith('_'):
            run = p.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = p.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        else:
            p.add_run(part)

