#!/usr/bin/env python3
"""
Mammoet 시트 타이틀 블록 양식 — PDF p.5 레이아웃·스타일 정확 재현.
순서: 상단(좌 Mammoet 주소 | 우 CALCULATION+문서번호) → Sheet 블록 → Client/Project/Subject → 하단 Rev~Checked 행 → 면책 → Condition.
"""

from docx import Document
from docx.shared import Mm, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge, value in kwargs.items():
        el = OxmlElement(f"w:{edge}")
        for k, v in value.items():
            el.set(qn(f"w:{k}"), str(v))
        tcPr.append(el)


def add_border_bottom(cell, pt=4):
    set_cell_border(
        cell, bottom={"val": "single", "sz": str(pt), "space": "0", "color": "000000"}
    )


def _run(p, text, bold=False, size_pt=None, font_name="Arial"):
    r = p.add_run(text)
    r.bold = bold
    if size_pt is not None:
        r.font.size = Pt(size_pt)
    r.font.name = font_name
    return r


def _para_8pt(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    for r in p.runs:
        r.font.size = Pt(8)
        r.font.name = "Arial"
    return p


def build_mammoet_sheet(doc, filled: bool) -> None:
    """PDF p.5 순서·스타일로 블록 구성. filled=True면 실제 값, False면 {{PLACEHOLDER}}."""
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(8)

    # ----- 1. 상단: 좌(Office + Mammoet 주소) | 우(CALCULATION + 문서번호 + Sap/Doc/Status) -----
    t_top = doc.add_table(rows=1, cols=2)
    t_top.columns[0].width = Mm(100)
    t_top.columns[1].width = Mm(90)
    t_top.allow_autofit = False
    left_cell = t_top.rows[0].cells[0]
    right_cell = t_top.rows[0].cells[1]

    # Left: Office + Mammoet Europe B.V. + address
    p = left_cell.paragraphs[0]
    p.clear()
    _run(p, "Office", size_pt=8)
    p = left_cell.add_paragraph()
    _run(p, "Mammoet Europe B.V.", size_pt=8)
    for line in [
        "Karel Doormanweg 47, Haven 580",
        "3115 JD Schiedam",
        "P.O. Box 570",
        "3100 AN Schiedam",
        "The Netherlands",
        "Phone +31 (0) 10 2042 424",
        "Fax +31 (0) 10 2042 442",
        "Website www.mammoet.com",
    ]:
        p = left_cell.add_paragraph()
        _run(p, line, size_pt=8)
    for p in left_cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    # Right: CALCULATION (20pt) + doc number + revision + For Approval; Sap Nr. / Doc. Nr. / Status
    p = right_cell.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, "CALCULATION", size_pt=20, bold=True)
    for line in [
        "4000266270" if filled else "{{DOCUMENT_NUMBER}}",
        "15111578-E-C16-00" if filled else "{{REVISION}}",
        "For Approval",
    ]:
        p = right_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _run(p, line, size_pt=8)
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, "Sap Nr." if not filled else "Sap Nr.", size_pt=8)
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, "Doc. Nr." if not filled else "Doc. Nr.", size_pt=8)
    p = right_cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _run(p, "Status", size_pt=8)
    for p in right_cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()

    # ----- 2. Sheet info 블록: Sheet info / Sheet name / MIND ID / Sheet revision -----
    t_sheet = doc.add_table(rows=4, cols=2)
    t_sheet.columns[0].width = Mm(50)
    t_sheet.columns[1].width = Mm(140)
    sheet_data = [
        ("Sheet info:", "" if filled else "{{SHEET_INFO}}"),
        ("Sheet name:", "12m RoRo Ramp" if filled else "{{SHEET_NAME}}"),
        ("MIND ID:", "MIND0-1868024057-3636" if filled else "{{MIND_ID}}"),
        ("Sheet revision:", "2.0" if filled else "{{SHEET_REVISION}}"),
    ]
    for i, (label, val) in enumerate(sheet_data):
        r = t_sheet.rows[i]
        r.cells[0].text = label
        r.cells[1].text = val
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(8)
                    run.font.name = "Arial"
        add_border_bottom(r.cells[1], pt=4)

    doc.add_paragraph()

    # ----- 3. Client / Project / Subject (라벨 8pt, 값 10pt) -----
    t_client = doc.add_table(rows=3, cols=2)
    t_client.columns[0].width = Mm(50)
    t_client.columns[1].width = Mm(140)
    client_data = [
        ("Client:", "Samsung C&T" if filled else "{{CLIENT}}"),
        ("Project:", "HVDC Lightning Al Ghallan Island" if filled else "{{PROJECT}}"),
        ("Subject:", "Check of 12m RoRo Ramp" if filled else "{{SUBJECT}}"),
    ]
    for i, (label, val) in enumerate(client_data):
        r = t_client.rows[i]
        r.cells[0].paragraphs[0].clear()
        _run(r.cells[0].paragraphs[0], label, size_pt=8)
        r.cells[1].paragraphs[0].clear()
        _run(r.cells[1].paragraphs[0], val, size_pt=10)
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
        add_border_bottom(r.cells[1], pt=4)

    doc.add_paragraph()

    # ----- 4. 하단: Rev. | Description | Date | Ref. | Checked (한 행, 9pt) -----
    t_rev = doc.add_table(rows=2, cols=5)
    for c in t_rev.columns:
        c.width = Mm(35)
    labels = ["Rev.:", "Description:", "Date:", "Ref.:", "Checked:"]
    values_filled = ["00", "For Approval", "08-12-2025", "SAhm  EMan", "971615  971580"]
    values_place = ["{{REV}}", "{{DESCRIPTION}}", "{{DATE}}", "{{REF}}", "{{CHECKED}}"]
    for ci, lbl in enumerate(labels):
        r = t_rev.rows[0].cells[ci]
        r.paragraphs[0].clear()
        _run(r.paragraphs[0], lbl, size_pt=9)
        r.paragraphs[0].paragraph_format.space_before = Pt(0)
        r.paragraphs[0].paragraph_format.space_after = Pt(0)
    for ci in range(5):
        r = t_rev.rows[1].cells[ci]
        r.paragraphs[0].clear()
        _run(
            r.paragraphs[0],
            values_filled[ci] if filled else values_place[ci],
            size_pt=8,
        )
        r.paragraphs[0].paragraph_format.space_before = Pt(0)
        r.paragraphs[0].paragraph_format.space_after = Pt(0)
        add_border_bottom(r, pt=4)

    doc.add_paragraph()

    # ----- 5. 면책 문구 (8pt) -----
    p = doc.add_paragraph()
    _run(
        p,
        "Without unauthorized signatures this document is uncontrolled, not binding and for indicative purpose only.",
        size_pt=8,
    )
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    # ----- 6. Condition + Terms (7pt) -----
    p = doc.add_paragraph()
    _run(p, "Condition", size_pt=7)
    p = doc.add_paragraph()
    _run(
        p,
        "Our Terms and Conditions apply to all our offers and agreements and any commitments rising therefrom.",
        size_pt=7,
    )
    p = doc.add_paragraph()
    _run(
        p,
        "We expressly deny the applicability of any other terms and conditions.",
        size_pt=7,
    )


def create_mammoet_sheet_template(out_path: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    build_mammoet_sheet(doc, filled=False)
    doc.save(out_path)
    print("Saved:", out_path)


def create_mammoet_sheet_filled(out_path: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    build_mammoet_sheet(doc, filled=True)
    doc.save(out_path)
    print("Saved:", out_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) > 1 and sys.argv[1] == "--filled":
        create_mammoet_sheet_filled("Mammoet_Sheet_Title_Block_Rev00.docx")
    else:
        out = (
            sys.argv[1]
            if len(sys.argv) > 1
            else "Mammoet_Sheet_Title_Block_Template.docx"
        )
        create_mammoet_sheet_template(out)
